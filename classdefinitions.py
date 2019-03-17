import datetime
import docx
import pickle
import traceback

now = datetime.datetime.now()

#The ultimate root referencer
class Base(object):
    def __init__(self):
        self.root = self
        
    def get_root(self):
        return self.root
    
    def get_parent(self):
        return self.parent

class Personal_Details(Base):
    def __init__(self):
        #set name from variable name. http://stackoverflow.com/questions/1690400/getting-an-instance-name-inside-class-init
        (filename,line_number,function_name,text)=traceback.extract_stack()[-2]
        def_name = text[:text.find('=')].strip()
        self.name = def_name
        
        self.FullName = 'Name'
        self.MobileNo = 'MobileNo'
        self.PersonalEmail = 'PersonalEmail'
        self.InternshipOrJob = 'InternshipOrJob'
        self.AvailableTimePeriod = 'AvailableTimePeriod'

class Current_Date(Base):
    def __init__(self):
        self.year = now.year
        self.month = now.month
        self.day = now.day

class Cover_Letter(Base):
    def __init__(self):
        self.SkillsPara1 = 'Relevant Skills Paragraph 1'
        self.SkillsPara2 = 'Relevant Skills Paragraph 2'
        self.WhyThisCompany = 'Even so, I wish to push the boundaries of technology, and use cutting edge technology to improve people’s lives. I acknowledge I do not have the resources myself to make this dream come true, but I am excited at the prospect of doing so with the highly qualified team at Entropica Labs.'
        self.ClosingParagraph = 'I believe that my self-driven attitude, my ability to work in a team, and my passion of hardware and software development will be beneficial to Entropica Labs, and I do hope you consider me for an internship position.'

class Additional_Information(Base):
    def __init__(self):
        self.Information = []
        for i in range(0,8):
            self.Information.append('Additional_Information' + str(i))
    pass

#class Overarching_Theme(Base):
#    def __init__(self):
#        (filename,line_number,function_name,text)=traceback.extract_stack()[-2]
#        def_name = text[:text.find('=')].strip()
#        self.OverarchingTheme = def_name
#        self.IndividualProjects = []
#    
#    def add_individual_project(self,IndividualProject):
#        self.IndividualProjects.append(IndividualProject)
        
class Individual_Project(Base):
    def __init__(self):
        self.OverarchingTheme = 'Overarching Theme'
        self.CompanyWorked = 'Place experience was attained eg. SUTD'
        self.Country = 'Country where experience was gained'
        self.JobScope = 'What did you do, eg. Student'
        self.TimePeriodWorked = 'Period Spent Working in the Job'
        self.ExperiencePoints = ['Got 3.1 GPA','something else','something else', 'something else', 'something else']

class Resume(Base):
    def __init__(self):
        self.overarching_theme = [Overarching_Theme()]
        self.additional_information = ['point1','point2','point3','point4','point5']

class ResumeDocument(Base):
    def __init__(self):
        self.document = docx.Document('Resume Template.docx')
        for para in self.document.paragraphs:
            for run in para.runs:
                #print(run.text)
                run.text = run.text.format(Name = self.get_root().personal_details.Name, 
                                           MobileNo = self.get_root().personal_details.MobileNo, 
                                           PersonalEmail = self.get_root().personal_details.PersonalEmail, 
                                           LinkedIn = '')
        
    def getText(self):
        doc = self.document
        fullText = []
        for para in doc.paragraphs:
            para = para.text
            fullText.append(para)
        return '\n'.join(fullText)
    
    def delete_paragraph(self):
        pass
    
    def document_combiner(self,source_document):
        if self == source_document:
            print('You cannot combine a document with itself')
            return
        for element in source_document.element.body:
            self.document.element.body.append(element)
        pass
    
    def individual_project_adder(self,IndividualProjectType):
        Output = docx.Document('Individual Projects Template.docx')
        for para in Output.paragraphs:
            #print(para.text)
            for run in para.runs:
                #print(run.text)
                run.text = run.text.format(JobScope = IndividualProjectType.JobScope,
                                           Country = IndividualProjectType.Country,
                                           TimePeriodWorked = IndividualProjectType.TimePeriodWorked,
                                           CompanyWorked = IndividualProjectType.CompanyWorked,
                                           ExperiencePoints = IndividualProjectType.ExperiencePoints)
                if run.text == '':
                    #Delete Paragraph Code
                    p = para._element
                    p.getparent().remove(p)
                    p._p = p.element = None
        self.document_combiner(Output)
        
    def overarching_theme_adder(self,OverarchingThemeType):
        Output = docx.Document('Overarching Theme Template.docx')
        for para in Output.paragraphs:
            #print(para.text)
            for run in para.runs:
                #print(run.text)
                run.text = run.text.format(OverarchingTheme = OverarchingThemeType.OverarchingTheme)
        self.document_combiner(Output)
        
    def additional_information_adder(self,AdditionalInformation):
        Output = docx.Document('Additional Information Template.docx')
        for para in Output.paragraphs:
            #print(para.text)
            for run in para.runs:
                #print(run.text)
                run.text = run.text.format(AdditionalInformation = AdditionalInformation)
                if run.text == '':
                    #Delete Paragraph Code
                    p = para._element
                    p.getparent().remove(p)
                    p._p = p.element = None
        self.document_combiner(Output)
        
    def output_resume(self):
        self.document.save('Resume Output.docx')
    

class Company_Application(Base):
    def __init__(self, CompanyName):
        self.CompanyName = CompanyName
        self.cover_letter = Cover_Letter()
        self.resume = Resume()
        self.resume_document = ResumeDocument()
        


#The Class File to end all classes
class Master_Class(Base):
    def __init__(self):
        
        #set name from variable name. http://stackoverflow.com/questions/1690400/getting-an-instance-name-inside-class-init
        (filename,line_number,function_name,text)=traceback.extract_stack()[-2]
        def_name = text[:text.find('=')].strip()
        self.name = def_name
#
#        try:
#            self.load()
#        
#        except:
#            print('error loading, now creating new')
        Base.root = self
        self.personal_details = Personal_Details()
        self.current_date = Current_Date()
        self.company_applications = {}
        #Additional info, skills para, overarchging themes and individual projects need to be implemented
        self.additional_information = []
        self.skills_paragraphs = []
        self.overarching_theme = []
        print('defaults loaded')

#            self.save()
        
    def add_company(self,CompanyName):
        if CompanyName not in self.company_applications:
            Name = CompanyName
            Company = Company_Application(CompanyName)
            self.company_applications[Name] = Company
        else:
            pass
            
#    def save(self):
#        print('saving' + self.name)
#        #save class as self.name.pickle
#        file = open(self.name+'.pickle','wb')
#        pickle.dump(self.__dir__(),file)
#        file.close()           
#        
#    def load(self):
#        
#        #try load self.name.txt
#        file = open(self.name+'.pickle','rb')      
#        self.__dir__ = pickle.load(file)
#        file.close()     
#        print('loaded' + self.name)

        
#def save_individual(class_obj):
#    print('saving ' + class_obj.name)
#    #save class as self.name.pickle
#    file = open(class_obj.name+'.pickle','wb')
#    pickle.dump(class_obj,file)
#    file.close()           
#    
#def load_individual(file_name_without_extension):
#    
#    #try load self.name.txt
#    file = open(file_name_without_extension+'.pickle','rb')      
#    class_obj = pickle.load(file)
#    file.close()     
#    print('loaded ' + class_obj.name)
#    Base.root = class_obj
#    return class_obj
    
#Shang = load_individual('Shang')

#Shang = Master_Class()
#print(Shang.__dir__)
#save_individual(Shang)

#
#OverarchingTheme1 = Overarching_Theme()
#IndividualProject1 = Individual_Project()
#AdditionalInformation = Additional_Information()
##
##
#Shang.add_company('gay')
##print(Shan.company_applications['gay'].resume.overarching_theme[0].Country)
##Shan.company_applications['gay'].resume_document.document_combiner(docx.Document('Additional Information Template.docx'))
#Shang.company_applications['gay'].resume_document.overarching_theme_adder(OverarchingTheme1)
#Shang.company_applications['gay'].resume_document.individual_project_adder(IndividualProject1)
#Shang.company_applications['gay'].resume_document.additional_information_adder(AdditionalInformation.Information)
#Shang.company_applications['gay'].resume_document.output_resume()
#
#save_individual(Shang)
#
#print(Shang.company_applications['gay'].resume_document.getText())
