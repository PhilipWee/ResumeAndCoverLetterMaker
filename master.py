import docx
#Necessary Details to put in the masterclass
Name = 'Philip Andrew Wee De Wang'
MobileNo = '84381245'
PersonalEmail = 'philip_andrew@mymail.sutd.edu.sg'
CompanyWorked = 'Singapore University of Technology and Education'
JobScope = 'Student'
Country = 'Singapore'
TimePeriodWorked = 'Jan 16 to Mar 16'
OverarchingTheme = 'EDUCATION'
ExperiencePoints = ['well thats pretty gay','really','gay','gay','']
AdditionalInformation = ['sexy','hot','chio','doctors hate me click here to find out how','','','','']

#Create a new Resume with the Resume Template
#implemented
Resume = docx.Document('Resume Template.docx')
print(getText(Resume))

#Used to get the raw text from a word document
#implemented
def getText(filename):
    doc = filename
    fullText = []
    for para in doc.paragraphs:
        para = para.text
    return '\n'.join(fullText)

#Function to delete paragraphs, necessary removing extra experience points in overarching theme maker
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
    
#Function to combine multiple documents, so we can spam fill the Resume document    
#implemented
def document_combiner(source_document,target_document):
    for element in source_document.element.body:
        target_document.element.body.append(element)
    return target_document
        
#Function to make the Overarching theme, Education, SUTD, School, etc
def OverarchingThemeMaker(OverarchingTheme,
                          CompanyWorked,
                          Country,
                          JobScope,
                          TimePeriodWorked,
                          ExperiencePoints):
    Output = docx.Document('Overarching Theme Template.docx')
    for para in Output.paragraphs:
        print(para.text)
        for run in para.runs:
#            print(run.text)
            run.text = run.text.format(OverarchingTheme = OverarchingTheme)
    return Output

def IndividualProjectMaker(OverarchingTheme,
                          CompanyWorked,
                          Country,
                          JobScope,
                          TimePeriodWorked,
                          ExperiencePoints):
    Output = docx.Document('Individual Projects Template.docx')
    for para in Output.paragraphs:
        print(para.text)
        for run in para.runs:
#            print(run.text)
            run.text = run.text.format(OverarchingTheme = OverarchingTheme,
                                       JobScope = JobScope,
                                       Country = Country,
                                       TimePeriodWorked = TimePeriodWorked,
                                       CompanyWorked = CompanyWorked,
                                       ExperiencePoints = ExperiencePoints)
            if run.text == '':
                delete_paragraph(para)
    return Output

def AdditionalInformationMaker(AdditionalInformation):
    Output = docx.Document('Additional Information Template.docx')
    for para in Output.paragraphs:
        print(para.text)
        for run in para.runs:
            print(run.text)
            run.text = run.text.format(AdditionalInformation = AdditionalInformation)
            if run.text == '':
                delete_paragraph(para)
    return Output

#The function to create the Resume
def create_resume(Resume):
    for para in Resume.paragraphs:
        for run in para.runs:
            print(run.text)
            run.text = run.text.format(Name = Name, 
                                       MobileNo = MobileNo, 
                                       PersonalEmail = PersonalEmail, 
                                       LinkedIn = '')
    
    Theme1 = OverarchingThemeMaker(OverarchingTheme,
                                      CompanyWorked,
                                      Country,
                                      JobScope,
                                      TimePeriodWorked,
                                      ExperiencePoints)
    
    IndividualProject1 = IndividualProjectMaker(OverarchingTheme,
                                              CompanyWorked,
                                              Country,
                                              JobScope,
                                              TimePeriodWorked,
                                              ExperiencePoints)
    
    IndividualProject2 = IndividualProjectMaker(OverarchingTheme,
                                              CompanyWorked,
                                              Country,
                                              JobScope,
                                              TimePeriodWorked,
                                              ExperiencePoints)
    
    AdditionalInformation1 = AdditionalInformationMaker(AdditionalInformation)

    Resume = document_combiner(Theme1, Resume)
    Resume = document_combiner(IndividualProject1, Resume)
    Resume = document_combiner(IndividualProject2, Resume)
    Resume = document_combiner(AdditionalInformation1, Resume)
    
    
    return Resume

create_resume(Resume)    
Resume.save('Resume Output.docx')


    

#OverarchingThemeMaker(OverarchingTheme,
#                                  CompanyWorked,
#                                  Country,
#                                  JobScope,
#                                  TimePeriodWorked,
#                                  ExperiencePoints).save('Resume Output.docx')


